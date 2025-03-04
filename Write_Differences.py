import os
import shutil
import re
from tkinter import Tk, filedialog, messagebox, Label, Button, Entry
from collections import Counter
from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import PyPDF2

# Common stop words (expand if needed)
STOP_WORDS = {
    "the", "and", "a", "an", "of", "in", "on", "at", "to", "for", "with", 
    "is", "it", "by", "this", "that", "from", "as", "but", "or"
}

# Clean a word to remove control characters or null bytes
def clean_word(word):
    return re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F]', '', word)

# Read text from different file formats
def extract_text(file_path):
    ext = os.path.splitext(file_path)[1].lower()

    if ext == '.docx':
        doc = Document(file_path)
        return [clean_word(word) for para in doc.paragraphs for word in para.text.split()]

    elif ext == '.pdf':
        words = []
        with open(file_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    words.extend(clean_word(word) for word in text.split())
        return words

    elif ext == '.txt':
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
            return [clean_word(word) for line in file for word in line.split()]

    else:
        raise ValueError(f"Unsupported file type: {ext}")

# Normalize words for matching (lowercase + strip punctuation)
def normalize_word(word):
    return word.lower().strip(".,!?()[]{}\"'")

# Highlight common words in DOCX
def highlight_text_in_docx(docx_path, common_words, output_path):
    doc = Document(docx_path)

    for para in doc.paragraphs:
        words = para.text.split()
        para.clear()

        for word in words:
            run = para.add_run(word + " ")
            if normalize_word(word) in common_words:
                highlight = parse_xml(r'<w:highlight {} w:val="yellow"/>'.format(nsdecls('w')))
                run._element.get_or_add_rPr().append(highlight)

    doc.save(output_path)

# Convert PDF or TXT to DOCX
def convert_to_docx(input_path, output_path):
    ext = os.path.splitext(input_path)[1].lower()

    doc = Document()

    if ext == '.pdf':
        with open(input_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    doc.add_paragraph(clean_word(text))

    elif ext == '.txt':
        with open(input_path, 'r', encoding='utf-8', errors='ignore') as file:
            for line in file:
                doc.add_paragraph(clean_word(line.strip()))

    doc.save(output_path)

# Save report listing matched words and counts
def save_comparison_report(output_folder, common_words_counter):
    report_path = os.path.join(output_folder, "comparison_report.txt")
    with open(report_path, 'w', encoding='utf-8') as report_file:
        report_file.write("Matched Words Report\n")
        report_file.write("=" * 40 + "\n")
        for word, count in common_words_counter.most_common():
            report_file.write(f"{word}: {count}\n")
    return report_path

# Core comparison logic
def run_comparison(file1, file2, output_folder):
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)

    # Extract, clean, normalize, and filter
    words1 = extract_text(file1)
    words2 = extract_text(file2)

    normalized1 = [normalize_word(word) for word in words1 if normalize_word(word) not in STOP_WORDS]
    normalized2 = [normalize_word(word) for word in words2 if normalize_word(word) not in STOP_WORDS]

    # Find common words
    common_words = set(normalized1).intersection(set(normalized2))

    # Count occurrences for the report
    common_words_counter = Counter(word for word in normalized1 if word in common_words)

    # Prepare output filename
    base_name = os.path.splitext(os.path.basename(file1))[0]
    compared_file = os.path.join(output_folder, f"{base_name}-compared.docx")

    # Convert file1 to DOCX if needed
    if file1.endswith('.docx'):
        shutil.copyfile(file1, compared_file)
    else:
        convert_to_docx(file1, compared_file)

    # Highlight common words in DOCX
    highlight_text_in_docx(compared_file, common_words, compared_file)

    # Save the comparison report
    report_file = save_comparison_report(output_folder, common_words_counter)

    return compared_file, report_file

# ---------------- GUI Section ----------------

def select_file1():
    file1_entry.delete(0, 'end')
    file_path = filedialog.askopenfilename(filetypes=[("Supported Files", "*.docx;*.pdf;*.txt")])
    file1_entry.insert(0, file_path)

def select_file2():
    file2_entry.delete(0, 'end')
    file_path = filedialog.askopenfilename(filetypes=[("Supported Files", "*.docx;*.pdf;*.txt")])
    file2_entry.insert(0, file_path)

def select_output_folder():
    output_entry.delete(0, 'end')
    folder = filedialog.askdirectory()
    output_entry.insert(0, folder)

def run():
    file1 = file1_entry.get().strip()
    file2 = file2_entry.get().strip()
    output_folder = output_entry.get().strip()

    if not (file1 and file2 and output_folder):
        messagebox.showerror("Error", "Please select both files and the output folder.")
        return

    try:
        compared_file, report_file = run_comparison(file1, file2, output_folder)
        messagebox.showinfo("Success", f"Comparison complete!\n\nHighlighted DOCX:\n{compared_file}\n\nReport:\n{report_file}")
    except Exception as e:
        messagebox.showerror("Error", f"An error occurred:\n{e}")

# ---------------- GUI Layout ----------------

root = Tk()
root.title("Compare and Highlight Tool V0.2")

Label(root, text="File 1 (DOCX, PDF, TXT)").grid(row=0, column=0, sticky="w", padx=5, pady=5)
file1_entry = Entry(root, width=50)
file1_entry.grid(row=0, column=1, padx=5, pady=5)
Button(root, text="Browse", command=select_file1).grid(row=0, column=2, padx=5, pady=5)

Label(root, text="File 2 (DOCX, PDF, TXT)").grid(row=1, column=0, sticky="w", padx=5, pady=5)
file2_entry = Entry(root, width=50)
file2_entry.grid(row=1, column=1, padx=5, pady=5)
Button(root, text="Browse", command=select_file2).grid(row=1, column=2, padx=5, pady=5)

Label(root, text="Output Folder").grid(row=2, column=0, sticky="w", padx=5, pady=5)
output_entry = Entry(root, width=50)
output_entry.grid(row=2, column=1, padx=5, pady=5)
Button(root, text="Browse", command=select_output_folder).grid(row=2, column=2, padx=5, pady=5)

Button(root, text="Run Comparison", command=run, bg="green", fg="white").grid(row=3, column=1, pady=20)

root.mainloop()
